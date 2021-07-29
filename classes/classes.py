import os
from docx import Document
import csv
from datetime import datetime
import docx2pdf

import docx


class App(object):
    CSV_FIELDS = ['correo', 'nombre', 'apellido', 'dni', 'pdf']
    
    def __init__(self, dest_path:str, word_template_path:str, csv_path:str, keywords_map:list) -> None:
        """
        Clase entorno para aplicacion de creacion de certificados
        
        :param dest_path: ubicacion de la carpeta de destino para guardar los certificados
        :param word_template: ubicacion del archivo word a utilizar como plantilla
        :param csv_path: ubicacion del archivo csv de donde se extraera informacion para certificados
        """
        self.destPath = dest_path
        self.template_path = word_template_path
        self.csv1_path = csv_path
        self.keywordsMap = keywords_map
        
    def create_certificates(self):
        """
        Metodo principal para iniciar creacion de certificados.
        """
        self.csv_partida = CsvLoader(self.csv1_path, 'r', self.CSV_FIELDS)
        self.csv_resultado = CsvLoader(os.path.join(self.destPath, 'res.csv'), 'w', self.CSV_FIELDS)
        persons = self.csv_partida.getContentAsList(firstLineHeaders=True)
        self.template = WordManage(self, self.template_path, self.keywordsMap)
        
        for correo, nombre, apellido, dni, pdf in persons:
            self.template.open_template()
            _person = Person(self, correo, nombre, apellido, dni)
            print(_person.__str__())
            _person.make_my_certificate()
            self.template.close_template()
        
        self.csv_partida.close()
        self.csv_resultado.close()

    def export_to_pdf(self):
        docx2pdf.convert(self.destPath)

    
class WordManage:
    def __init__(self, app: App, word_path:str, keywords:list) -> None:
        """
        Clase para manipulacion rapida de archivos word
        
        :param word_path: ubicacion del archivo word 
        :param keywords: lista con las palabras claves a remplazar
        """
        self.app = app
        self.path = word_path
        self.keywords = keywords
        
    def open_template(self):
        self.docx = Document(self.path)
        
    def close_template(self):
        del self.docx
        
    def replace(self, info:list):
        """
        Metodo para reemplazo de palabras claves por valores.
        
        :param info: lista con la informacion que se reemplazara
        """
        if len(info) == len(self.keywords):
            for keyword, item in zip(self.keywords, info):
                try:
                    for table in self.docx.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    for run in p.runs:
                                        if keyword in run.text:
                                            run.text = run.text.replace(keyword, item)
                                            print(f'Se reemplazo {item} en key {keyword}')
                except:
                    print(f'Error al remplazar informacion') 
                    raise
                    
        else:
            print('Deben haber iguales cantidades de argumentos y de palabras claves')
    
    def save_as(self, _dest_path:str, file_name:str):
        try:
            self.docx.save(f'{_dest_path}\\{file_name.upper()}.docx')
        except:
            print(f'Error al guardar')


class Person:
    def __init__(self, _app:App, _correo:str, _nombre:str, _apellido:str, _dni:str):
        self.app = _app
        self.correo = _correo
        self.nombre = _nombre.upper()
        self.apellido = _apellido.upper()
        self.dni = self.decore_dni(_dni)
    
    def __str__(self):
        return f'\n{self.nombre} {self.apellido}, dni: {self.dni}, correo: {self.correo}'
    
    def get_full_name(self) -> str:
        return f'{self.nombre} {self.apellido}'
    
    def make_my_certificate(self):
        print(f'Valores a reemplazar: {self.app.keywordsMap}')
        
        info = [self.get_full_name(), self.dni]
        
        print(f'Valores que se reemplazaran: {info}')
        
        self.app.template.replace(info)
        
        self.app.template.save_as(self.app.destPath, f'{self.get_full_name()}.docx')
        
        self.app.csv_resultado.write_row([self.correo, self.nombre, self.apellido, self.clear_dni(self.dni),
                                          os.path.join(self.app.destPath, f'{self.get_full_name().upper()}.pdf')])
        
    def clear_dni(self, dni:str) -> str:
        """
        Esta funcion corrige el formato del dato dni cuando sea necesario. Devuelve el dato eliminando cualquier aparicion de espacios
        y puntos.
        
        :param dni: str con el dato de la persona
        :return res: str con el dato formateado
        """
        res = dni.replace(' ', '').replace('.', '')
        return res
        
    def decore_dni(self, _dni:str):
        dni = _dni
        if '.' not in dni:
            dni = f'{dni[:2]}.{dni[2:5]}.{dni[5:]}'
            return dni
        else:
            dni.replace('.', '')
            self.fix(dni)
            

class CsvLoader(object):
    def __init__(self, csvPath: str, mode: str, fields: list):
        """
        Constructor para la clase CSVLoader.
        :param csvPath: str path del archivo
        :param mode: str r(read)/w(write)/a(append)
        PRECAUCION: w mode sobreescribira un archivo como vacio
        :param fields:
        """

        self.csvFile = open(csvPath, mode=mode, encoding='UTF-8', newline='')
        self.fields = fields

        if self.csvFile.mode == 'r':
            self.reader = csv.DictReader(self.csvFile, fields)
        elif self.csvFile.mode == 'w' or self.csvFile.mode == 'a':
            self.writer = csv.DictWriter(self.csvFile, fields)
            if self.csvFile.mode == 'w':
                self.writer.writeheader()

    def getContentAsList(self, firstLineHeaders=False):
        """
        En caso de que el modo sea r, se crea un reader, con el cual se puede obtener todo el contenido del archivo como
        una lista.
        :param firstLineHeaders: bool
        :return: list
        """
        try:
            contentList: list = []
            row: list = []
            for i in self.reader:
                for field in self.fields:
                    row.append(i[field])
                contentList.append(row)
                row = []
            if firstLineHeaders:
                return contentList[1:]
            else:
                return contentList
        except AttributeError as e:
            print(f'there is no reader.\nERROR: {e}')
            return None

    def write_row(self, row: list):
        try:
            dict_row = {}
            
            for key, value in zip(self.fields, row):
                dict_row[key] = value
                
            self.writer.writerow(dict_row)
        except AttributeError as e:
            print(f'there is no writer.\nERROR: {e}')
        
    def close(self):
        """
        Funcion para cerrar archivo leido.
        :return:
        """
        self.csvFile.close()

