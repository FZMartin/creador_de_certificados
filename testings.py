import datetime
import docx
from docx.shared import Pt
from datetime import datetime

def main():
    path = 'C:\\Users\\pedro\\desktop\\certificado_taller1.docx'
    destpath = 'C:\\Users\\pedro\\desktop\\PEDROZO JUAN MARTIN.docx'
    document = docx.Document(path)
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        print(run.text)
                        if run.text == "person_info":
                            run.text = 'PEDROZO JUAN MARTIN DNI 43.944.733'
                            break
    document.save(destpath)

def fix(_dni):
    dni = _dni
    if '.' not in dni:
        dni = f'{dni[:2]}.{dni[2:5]}.{dni[5:]}'
        return dni
    else:
        dni.replace('.', '')
        fix(dni)
    
    
if __name__ == '__main__':
    main()